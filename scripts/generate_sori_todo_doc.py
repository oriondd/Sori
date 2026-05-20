from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Sori_MVP_Page_By_Page_Todo.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    paragraph.paragraph_format.first_line_indent = Inches(-0.1)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        add_bullet(doc, "[ ] " + item)


def add_page_section(doc: Document, title: str, status: str, now: list[str], todo: list[str]) -> None:
    doc.add_heading(title, level=2)
    status_p = doc.add_paragraph()
    status_p.paragraph_format.space_after = Pt(4)
    status_run = status_p.add_run("Current status: ")
    status_run.bold = True
    status_p.add_run(status)

    doc.add_paragraph("What is already in place:", style="Heading 3")
    for item in now:
        add_bullet(doc, item)

    doc.add_paragraph("Still needs work:", style="Heading 3")
    add_checklist(doc, todo)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sori MVP Page-by-Page To-Do List")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Working checklist for the current local Expo/Supabase prototype").italic = True

    doc.add_heading("How To Use This Document", level=1)
    add_bullet(doc, "Use the Priority 1 section as the near-term build order.")
    add_bullet(doc, "Use each page section as a running checklist while testing localhost.")
    add_bullet(doc, "Items marked as architecture/security should be handled before public beta.")
    add_bullet(doc, "This document reflects the current repo state as of May 20, 2026.")

    doc.add_heading("Build Snapshot", level=1)
    snapshot = doc.add_table(rows=1, cols=3)
    snapshot.style = "Table Grid"
    hdr = snapshot.rows[0].cells
    for cell, text in zip(hdr, ["Area", "Current State", "Next Decision"]):
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, text, True)
    rows = [
        ("Frontend", "Expo Router / React Native Web prototype", "Keep web-first polish while preparing mobile-specific media and layout paths."),
        ("Auth", "Supabase Auth wired for signup/login/handle claim", "Move identity/profile data into durable profile tables with RLS."),
        ("Profiles", "Custom HTML/CSS/JS iframe sandbox with local safety bones", "Add production-grade moderation, storage, and public profile routes."),
        ("Feed", "Local browser post store with text/media, edit/delete, visibility labels", "Move posts/media/likes/comments to Supabase and enforce privacy."),
        ("Security", "Local sanitizer, sandbox guard, link warning, report queue placeholders", "Add tests, real moderation API, admin roles, and audit logs."),
    ]
    for row in rows:
        cells = snapshot.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    doc.add_heading("Priority 1: Next Build Order", level=1)
    priority_items = [
        "Replace localStorage post/profile persistence with Supabase tables and Storage buckets.",
        "Create database schema and RLS for posts, media, comments, likes, profiles, handles, reports, and tiers.",
        "Add route guards so public landing, authenticated feed, and profile pages have clear ownership.",
        "Make media upload production-shaped: file validation, thumbnails, compression, progress, retry, and delete cleanup.",
        "Build privacy enforcement for Public, Friends only, and Private posts at the query/policy layer.",
        "Add real profile public URLs such as /@sori or /profile/sori.",
        "Add automated security tests for iframe sandbox behavior and unsafe link blocking.",
        "Connect the AI Profile Stylist to a safe server route with strict JSON output and tier limits.",
    ]
    add_checklist(doc, priority_items)

    pages = [
        (
            "Public Landing / Splash",
            "Early visual landing page work exists conceptually, but the current root route is primarily the authenticated feed.",
            [
                "Sori visual identity and left rail navigation are established.",
                "Signup/login routes exist separately.",
            ],
            [
                "Split unauthenticated visitors to a real public landing page and logged-in users to the feed.",
                "Add a polished hero visualizer showing profile canvas, storefront grid, feed, marketplace, and AI stylist.",
                "Add clear signup/login CTAs and responsive mobile layout.",
                "Add SEO metadata and share preview images for web launch.",
            ],
        ),
        (
            "Signup / Login / Claim Handle",
            "Supabase Auth works locally, confirmation email works, and handle claim flow exists.",
            [
                "Signup, login, and claim-handle pages exist.",
                "Founder handle and founder badge behavior are prototyped.",
                "Identity badge displays handle and verified badge.",
            ],
            [
                "Create Supabase profile/handle tables with unique handle constraint.",
                "Move founder verification out of client metadata into admin-controlled database fields.",
                "Add password reset, resend confirmation, logout, and account settings flows.",
                "Customize Supabase email sender/template branding for Sori.",
                "Add auth loading boundaries for all protected pages.",
            ],
        ),
        (
            "Feed Page",
            "Feed can show locally saved posts with text/media, identity, visibility pill, edit/delete menu, and starter actions.",
            [
                "Global feed route exists at root.",
                "Post cards support edit, delete, and visibility changes locally.",
                "Text posts and local media rendering are prototyped.",
            ],
            [
                "Move posts to Supabase with durable media URLs.",
                "Implement comments, likes, repost/share, save/bookmark, and report post.",
                "Enforce visibility in backend queries and RLS.",
                "Add pagination/infinite scroll and loading skeletons.",
                "Add empty states for new accounts and follow suggestions.",
                "Add video thumbnail, duration, mute/autoplay controls, and mobile-safe playback.",
            ],
        ),
        (
            "Create Post",
            "Composer supports text, photo/video buttons, visibility dropdown, local previews, posting, and redirect to feed.",
            [
                "Public/Friends only/Private selector exists.",
                "Photo and video flows are separated.",
                "Post button redirects to feed after save.",
            ],
            [
                "Use Supabase Storage upload instead of local browser blob storage.",
                "Add upload progress, cancel, retry, and file size/type validation.",
                "Generate thumbnails and optimized image/video variants.",
                "Add mobile Expo media picker support.",
                "Add post draft autosave and discard confirmation.",
                "Add moderation pre-check before publish.",
            ],
        ),
        (
            "My Profile",
            "Profile page supports default profile, custom iframe profile, safe mode, report profile, Top 10 placeholder, and security delay.",
            [
                "Custom profiles render in sandboxed iframe.",
                "Report Profile modal and local hidden-pending-review behavior exist.",
                "Top 10 grid is visible as a compact profile element.",
            ],
            [
                "Create public profile URL and viewer/owner modes.",
                "Build Top 10 editor with friends/search/reorder/custom border styling.",
                "Add profile avatar/header media upload.",
                "Add profile posts/media tabs.",
                "Persist profile views, likes, pinned badges, and custom profile versions in Supabase.",
                "Improve loading transition so custom iframe never flashes default content.",
            ],
        ),
        (
            "Customize Profile",
            "Customizer supports free themes, custom HTML/CSS/JS, AI draft placeholder, preview iframe, save, reset, and version history.",
            [
                "HTML/CSS/JS editors exist as text areas.",
                "Preview uses sandboxed iframe.",
                "Version history/reset and size limits are now present.",
            ],
            [
                "Replace text areas with a real code editor experience.",
                "Add linting, formatting, and clear sanitizer warnings.",
                "Add AI Profile Stylist server endpoint with JSON-only output.",
                "Add tier gates for advanced JS, effects, marketplace themes, and premium layouts.",
                "Add theme import/export and preview thumbnails.",
                "Add performance monitor and auto safe-mode if profile hangs.",
            ],
        ),
        (
            "Designer Market",
            "Market page exists as an early placeholder route.",
            [
                "Navigation entry exists.",
            ],
            [
                "Build marketplace listing grid with previews, categories, search, and filters.",
                "Create theme detail page with live preview in sandbox.",
                "Add seller upload flow, pricing, approval status, and theme versioning.",
                "Connect Stripe Connect for marketplace payouts.",
                "Add purchase/install flow and license records.",
            ],
        ),
        (
            "Dashboard / Profile Studio",
            "Dashboard exists as a profile studio placeholder with identity strip and launch checklist.",
            [
                "Dashboard route exists.",
                "Basic checklist and stats cards are displayed.",
            ],
            [
                "Connect real profile stats: views, likes, grid clicks, post reach.",
                "Add tabs for Profile, Grid, Posts, Security, Billing, and Account.",
                "Add onboarding progress and next recommended action.",
                "Add tier status, subscription management, and feature unlocks.",
            ],
        ),
        (
            "Admin Review / Security Test",
            "Local security bones exist with admin review and sandbox test pages.",
            [
                "/security-test lists guardrail test targets.",
                "/admin-review shows local report queue placeholder.",
                "Sandbox blocks forms/popups/private links and routes external links through a warning modal.",
            ],
            [
                "Add automated tests for each listed sandbox/security behavior.",
                "Move admin review behind role-based access.",
                "Connect report queue to Supabase with audit logs.",
                "Integrate real image moderation provider.",
                "Add appeal/unhide workflow for profiles hidden by reports.",
            ],
        ),
        (
            "Subscriptions / Monetization",
            "Tier concepts and schema plans exist, but payment flows are not implemented in the app UI yet.",
            [
                "TierGate component/spec has been planned in earlier architecture.",
                "Stripe webhook/schema requirements are documented conceptually.",
            ],
            [
                "Create subscription tables and RLS migrations.",
                "Build Stripe checkout/session creation endpoint.",
                "Build Stripe webhook handler and price-to-tier mapping.",
                "Add billing page and tier upgrade UI.",
                "Gate premium features: 4D visualizer, audio/downloads, checkout links, business storefront.",
            ],
        ),
    ]

    doc.add_heading("Page-by-Page Checklist", level=1)
    for page in pages:
        add_page_section(doc, *page)

    doc.add_heading("Cross-Cutting Engineering Checklist", level=1)
    cross_cutting = [
        "Add unit/integration tests for auth, posts, profile security, sanitizer, and URL validation.",
        "Add lint script and CI workflow on GitHub.",
        "Create Supabase migrations folder and seed data for local development.",
        "Move secrets into environment templates and document setup in README.",
        "Add error boundary and app-wide loading transition layer.",
        "Add accessibility pass for keyboard navigation, labels, contrast, and focus states.",
        "Add responsive QA checklist for desktop, tablet, and phone widths.",
        "Add analytics events for signup, post creation, profile view, theme save, and purchase intent.",
    ]
    add_checklist(doc, cross_cutting)

    doc.add_heading("Definition of Done for the Next Milestone", level=1)
    done_items = [
        "A new user can sign up, confirm email, claim a unique handle, and land on the feed.",
        "A user can create text/photo/video posts that persist across refreshes and another browser session.",
        "A user can customize their profile, save it, reopen the editor, and revert a previous version.",
        "A visitor can view a profile safely without user code touching the parent app.",
        "External profile links always show the Sori warning before leaving.",
        "Reports land in an admin review queue.",
        "GitHub has clear setup instructions so another developer can run the project quickly.",
    ]
    add_checklist(doc, done_items)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Sori MVP checklist - local planning document").font.size = Pt(9)

    doc.save(OUT)


if __name__ == "__main__":
    build()
